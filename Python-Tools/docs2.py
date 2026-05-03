#!/usr/bin/env python3
"""
VAPT Report Word Document Generator
=====================================
Reads the VAPT-structured JSON from script.py and generates a Word document
with formatted vulnerability tables matching the standard VAPT report template.

Usage:
    python generate_docx.py [enriched_report.json] [output.docx]

Dependencies:
    pip install python-docx
"""

import json
import sys
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------- STYLING ----------
SEVERITY_COLORS = {
    'Critical': RGBColor(0xC0, 0x00, 0x00),    # #C00000
    'High':     RGBColor(0xEE, 0x00, 0x00),    # #EE0000
    'Medium':   RGBColor(0xFF, 0xC0, 0x00),    # #FFC000
    'Low':      RGBColor(0xFF, 0xFF, 0x00),    # #FFFF00
    'Info':     RGBColor(0x00, 0xB0, 0xF0),    # #00B0F0
}

SEVERITY_BG = {
    'Critical': 'C00000',
    'High':     'EE0000',
    'Medium':   'FFC000',
    'Low':      'FFFF00',
    'Info':     '00B0F0',
}

# Text color on severity badge (dark bg = white text, light bg = black text)
SEVERITY_TEXT = {
    'Critical': RGBColor(0xFF, 0xFF, 0xFF),
    'High':     RGBColor(0xFF, 0xFF, 0xFF),
    'Medium':   RGBColor(0x00, 0x00, 0x00),
    'Low':      RGBColor(0x00, 0x00, 0x00),
    'Info':     RGBColor(0xFF, 0xFF, 0xFF),
}

HEADER_BG = 'DAEEF3'  # #DAEEF3 — header / border shade
LABEL_BG  = 'DAEEF3'  # Same shade for label rows

# Severity sort order
SEVERITY_ORDER = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3, 'Info': 4}


def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_valign(cell, align=WD_ALIGN_VERTICAL.CENTER):
    """Set vertical alignment of a table cell."""
    cell.vertical_alignment = align


def set_cell_text(cell, text, bold=False, font_size=10, font_name='Calibri',
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Set text in a cell with formatting. Vertically centered, horizontally left."""
    cell.text = ''
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # Set spacing
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def set_cell_multiline(cell, lines, font_size=10, font_name='Calibri', bold=False):
    """Set multiple lines in a cell."""
    cell.text = ''
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(lines):
        if i > 0:
            run = p.add_run('\n')
            run.font.size = Pt(font_size)
        run = p.add_run(str(line))
        run.font.size = Pt(font_size)
        run.font.name = font_name
        run.bold = bold


def merge_row_cells(table, row_idx, start_col, end_col):
    """Merge cells in a row from start_col to end_col (inclusive)."""
    cell_a = table.cell(row_idx, start_col)
    cell_b = table.cell(row_idx, end_col)
    cell_a.merge(cell_b)


def add_label_row(table, row_idx, label, value, col_span_all=True):
    """Add a label header row followed by a value row."""
    # Label row
    row = table.rows[row_idx]
    if col_span_all:
        merge_row_cells(table, row_idx, 0, 2)
    cell = table.cell(row_idx, 0)
    set_cell_text(cell, label, bold=True, font_size=10)
    set_cell_bg(cell, LABEL_BG)


# ---------- VULNERABILITY TABLE BUILDER ----------
def create_vuln_table(doc, vuln):
    """Create a single vulnerability table matching the VAPT report format."""

    # Section heading
    heading = doc.add_heading(level=3)
    run = heading.add_run(f"{vuln['section_number']}  {vuln['title']}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # Create table: 3 columns, multiple rows
    table = doc.add_table(rows=20, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(7.5)
        row.cells[2].width = Cm(3.5)

    row_idx = 0

    # ── Row 0: Header labels ──
    set_cell_text(table.cell(row_idx, 0), 'Reference No:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), HEADER_BG)
    set_cell_text(table.cell(row_idx, 1), 'Risk Rating:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 1), HEADER_BG)
    set_cell_text(table.cell(row_idx, 2), 'Status', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 2), HEADER_BG)

    # ── Row 1: Header values ──
    row_idx += 1
    set_cell_text(table.cell(row_idx, 0), vuln['reference_no'], font_size=10)
    severity = vuln['risk_rating']
    set_cell_text(table.cell(row_idx, 1), severity, bold=True, font_size=10,
                  color=SEVERITY_TEXT.get(severity, RGBColor(0xFF, 0xFF, 0xFF)))
    set_cell_bg(table.cell(row_idx, 1), SEVERITY_BG.get(severity, '808080'))
    set_cell_text(table.cell(row_idx, 2), vuln['status'], font_size=10,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Row 2: Tools / Category labels ──
    row_idx += 1
    merge_row_cells(table, row_idx, 1, 2)
    set_cell_text(table.cell(row_idx, 0), 'Tools Used:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)
    set_cell_text(table.cell(row_idx, 1), 'Category:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 1), LABEL_BG)

    # ── Row 3: Tools / Category values ──
    row_idx += 1
    merge_row_cells(table, row_idx, 1, 2)
    set_cell_text(table.cell(row_idx, 0), vuln['tools_used'], font_size=10)
    set_cell_text(table.cell(row_idx, 1), vuln['category'], font_size=10)

    # ── Row 4: CWE / CVSS labels ──
    row_idx += 1
    merge_row_cells(table, row_idx, 1, 2)
    set_cell_text(table.cell(row_idx, 0), 'CWE:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)
    set_cell_text(table.cell(row_idx, 1), 'CVSS:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 1), LABEL_BG)

    # ── Row 5: CWE / CVSS values ──
    row_idx += 1
    merge_row_cells(table, row_idx, 1, 2)
    set_cell_text(table.cell(row_idx, 0), vuln['cwe_id'], font_size=10)
    set_cell_text(table.cell(row_idx, 1), vuln['cvss_score'], font_size=10)

    # ── Row 6: CVE label ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), 'CVE:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)

    # ── Row 7: CVE value ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), vuln['cve_id'], font_size=10)

    # ── Row 8: Description label ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), 'Vulnerability Description:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)

    # ── Row 9: Description value ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), vuln['description'], font_size=10)

    # ── Row 10: Identified by label ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0),
                  'Vulnerability Identified By / How It Was Discovered:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)

    # ── Row 11: Identified by value ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), vuln['identified_by'], font_size=10)

    # ── Row 12: Identifiers label ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), 'Identifiers:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)

    # ── Row 13: Identifiers value ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    idents = vuln.get('identifiers', [])
    set_cell_multiline(table.cell(row_idx, 0), idents if idents else ['N/A'], font_size=9)

    # ── Row 14: File Path label ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), 'File Path:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)

    # ── Row 15: File Path value ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    paths = vuln.get('file_paths', [])
    set_cell_multiline(table.cell(row_idx, 0), paths if paths else ['N/A'], font_size=9)

    # ── Row 16: Countermeasures label ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), 'Suggested Countermeasures:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)

    # ── Row 17: Countermeasures value ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), vuln['countermeasures'], font_size=10)

    # ── Row 18: References label ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    set_cell_text(table.cell(row_idx, 0), 'References:', bold=True, font_size=10)
    set_cell_bg(table.cell(row_idx, 0), LABEL_BG)

    # ── Row 19: References value ──
    row_idx += 1
    merge_row_cells(table, row_idx, 0, 2)
    refs = vuln.get('references', [])
    set_cell_multiline(table.cell(row_idx, 0), refs if refs else ['N/A'], font_size=9)

    # Add spacing after table
    doc.add_paragraph('')


# ---------- SUMMARY TABLE ----------
def create_summary_table(doc, meta, vulns):
    """Create a summary table at the top of the document."""
    doc.add_heading('Vulnerability Summary', level=2)

    # Severity summary
    summary = meta.get('severity_summary', {})
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Critical', 'High', 'Medium', 'Low', 'Info']
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, font_size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(table.cell(0, i), SEVERITY_BG.get(h, '808080'))
        set_cell_text(table.cell(1, i), str(summary.get(h, 0)), font_size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph('')

    # Vulnerability index table
    doc.add_heading('Vulnerability Index', level=2)
    idx_table = doc.add_table(rows=1 + len(vulns), cols=5)
    idx_table.style = 'Table Grid'
    idx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    idx_headers = ['#', 'Reference No', 'CVE ID', 'Vulnerability', 'Severity']
    for i, h in enumerate(idx_headers):
        set_cell_text(idx_table.cell(0, i), h, bold=True, font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(idx_table.cell(0, i), HEADER_BG)

    # Rows
    for row_num, v in enumerate(vulns, 1):
        set_cell_text(idx_table.cell(row_num, 0), str(row_num), font_size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(idx_table.cell(row_num, 1), v['reference_no'], font_size=9)
        set_cell_text(idx_table.cell(row_num, 2), v['cve_id'], font_size=9)
        set_cell_text(idx_table.cell(row_num, 3), v['title'], font_size=9)
        sev = v['risk_rating']
        set_cell_text(idx_table.cell(row_num, 4), sev, font_size=9, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      color=SEVERITY_COLORS.get(sev, RGBColor(0, 0, 0)))

    # Set column widths
    widths = [Cm(1), Cm(3.5), Cm(3.5), Cm(6), Cm(2.5)]
    for row in idx_table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.add_paragraph('')


# ---------- MAIN ----------
def main():
    json_file = sys.argv[1] if len(sys.argv) > 1 else 'enriched_report.json'
    docx_file = sys.argv[2] if len(sys.argv) > 2 else 'vapt_report.docx'

    if not os.path.exists(json_file):
        print(f"Error: {json_file} not found.")
        print("Run script.py first to generate the JSON file.")
        sys.exit(1)

    print(f"Reading {json_file}...")
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    meta = data.get('report_meta', {})
    vulns = data.get('vulnerabilities', [])

    if not vulns:
        print("No vulnerabilities found in JSON.")
        sys.exit(1)

    print(f"Found {len(vulns)} vulnerabilities. Generating Word document...")

    # Sort by severity: Critical → High → Medium → Low → Info
    vulns.sort(key=lambda v: SEVERITY_ORDER.get(v.get('risk_rating', 'Info'), 99))

    # Re-number sections after sorting
    for i, v in enumerate(vulns, 1):
        v['section_number'] = f"3.1.{i}"

    print(f"  Sorted: {sum(1 for v in vulns if v['risk_rating']=='Critical')} Critical, "
          f"{sum(1 for v in vulns if v['risk_rating']=='High')} High, "
          f"{sum(1 for v in vulns if v['risk_rating']=='Medium')} Medium, "
          f"{sum(1 for v in vulns if v['risk_rating']=='Low')} Low")

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Title page
    title = doc.add_heading('VAPT Report – SCA Findings', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph(f"Generated: {meta.get('generated', 'N/A')}")
    doc.add_paragraph(f"Total Packages: {meta.get('total_packages', 'N/A')}")
    doc.add_paragraph(f"Total Vulnerabilities: {meta.get('total_vulnerabilities', 'N/A')}")
    doc.add_paragraph(f"CVE Year Filter: ≥ {meta.get('cve_year_filter', 'N/A')}")
    doc.add_paragraph('')

    # Summary tables
    create_summary_table(doc, meta, vulns)

    # Section header
    doc.add_page_break()
    doc.add_heading('3.1  SCA Vulnerability Details', level=2)

    # Create a table for each vulnerability
    for i, vuln in enumerate(vulns):
        print(f"  [{i+1}/{len(vulns)}] {vuln['cve_id']} - {vuln['title'][:50]}...")
        create_vuln_table(doc, vuln)

    # Save
    doc.save(docx_file)
    print(f"\nDone! Saved {docx_file} with {len(vulns)} vulnerability tables.")


if __name__ == '__main__':
    main()
